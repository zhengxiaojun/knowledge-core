import io
import magic
import docx
from pypdf import PdfReader
from fastapi import UploadFile
import pandas as pd
from pathlib import Path

class UnsupportedContentTypeError(Exception):
    """Exception raised for unsupported file types."""
    def __init__(self, content_type: str):
        self.content_type = content_type
        super().__init__(f"Unsupported content type: {content_type}")

def _parse_docx(file_stream: io.BytesIO) -> str:
    """Parses a .docx file and returns its text content."""
    document = docx.Document(file_stream)
    return "\n".join([para.text for para in document.paragraphs])

def _parse_pdf(file_stream: io.BytesIO) -> str:
    """Parses a .pdf file and returns its text content."""
    reader = PdfReader(file_stream)
    return "\n".join([page.extract_text() for page in reader.pages])

async def parse_document(file: UploadFile) -> str:
    """
    Parses an uploaded file and extracts text content based on its MIME type.
    Supports plain text, DOCX, and PDF.
    """
    content = await file.read()
    
    # Reset stream position just in case
    await file.seek(0)
    
    # Use python-magic to determine the file type from its content
    mime_type = magic.from_buffer(content, mime=True)
    
    file_stream = io.BytesIO(content)

    if "text/plain" in mime_type:
        return content.decode("utf-8")
    elif "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in mime_type:
        return _parse_docx(file_stream)
    elif "application/pdf" in mime_type:
        return _parse_pdf(file_stream)
    else:
        raise UnsupportedContentTypeError(mime_type)


class DocumentParser:
    """
    Document parser for various file formats.
    Supports Word, PDF, Excel, and text files.
    """

    def parse_word(self, file_path: str) -> str:
        """Parse Word document (.docx) and extract text content."""
        try:
            document = docx.Document(file_path)
            paragraphs = [para.text for para in document.paragraphs if para.text.strip()]

            # Also extract text from tables
            tables_text = []
            for table in document.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        tables_text.append(" | ".join(row_text))

            all_text = paragraphs + tables_text
            return "\n".join(all_text)
        except Exception as e:
            raise Exception(f"Failed to parse Word document: {str(e)}")

    def parse_pdf(self, file_path: str) -> str:
        """Parse PDF document and extract text content."""
        try:
            reader = PdfReader(file_path)
            text_content = []

            for page in reader.pages:
                text = page.extract_text()
                if text.strip():
                    text_content.append(text)

            return "\n".join(text_content)
        except Exception as e:
            raise Exception(f"Failed to parse PDF document: {str(e)}")

    def parse_excel(self, file_path: str) -> str:
        """Parse Excel file and extract content as text."""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            content = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                # Add sheet name
                content.append(f"\n=== Sheet: {sheet_name} ===")

                # Convert dataframe to text representation
                # Remove rows that are all NaN
                df = df.dropna(how='all')

                # Convert to string representation
                if not df.empty:
                    # Get column headers
                    headers = " | ".join([str(col) for col in df.columns])
                    content.append(headers)
                    content.append("-" * len(headers))

                    # Get row data
                    for _, row in df.iterrows():
                        row_text = " | ".join([str(val) if pd.notna(val) else "" for val in row])
                        content.append(row_text)

            return "\n".join(content)
        except Exception as e:
            raise Exception(f"Failed to parse Excel file: {str(e)}")

    def parse_text(self, file_path: str) -> str:
        """Parse plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='gbk') as f:
                return f.read()
        except Exception as e:
            raise Exception(f"Failed to parse text file: {str(e)}")

    def parse(self, file_path: str, file_type: str = None) -> str:
        """
        Auto-detect and parse document based on file extension or specified type.

        Args:
            file_path: Path to the file
            file_type: Optional file type hint (word, pdf, excel, text)

        Returns:
            Extracted text content
        """
        path = Path(file_path)

        if file_type:
            type_map = {
                'word': self.parse_word,
                'doc': self.parse_word,
                'docx': self.parse_word,
                'pdf': self.parse_pdf,
                'excel': self.parse_excel,
                'xlsx': self.parse_excel,
                'xls': self.parse_excel,
                'text': self.parse_text,
                'txt': self.parse_text
            }
            parser_func = type_map.get(file_type.lower())
            if parser_func:
                return parser_func(file_path)

        # Auto-detect from extension
        ext = path.suffix.lower()
        if ext in ['.docx', '.doc']:
            return self.parse_word(file_path)
        elif ext == '.pdf':
            return self.parse_pdf(file_path)
        elif ext in ['.xlsx', '.xls']:
            return self.parse_excel(file_path)
        elif ext in ['.txt', '.md']:
            return self.parse_text(file_path)
        else:
            raise UnsupportedContentTypeError(f"Unsupported file extension: {ext}")

